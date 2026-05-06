# -*- coding: utf-8 -*-
"""Export assembled paper sections to downloadable files."""

from __future__ import annotations

import base64
import io
import os
import re
import shutil
import subprocess
import tempfile
from dataclasses import dataclass
from urllib.parse import unquote, urlparse


@dataclass
class ExportedPaper:
    filename: str
    content_type: str
    body: bytes


IMAGE_LINE_RE = re.compile(r'^\s*!\[([^\]]*)\]\(([^)]+)\)\s*$')
MARKDOWN_IMAGE_RE = re.compile(r'!\[([^\]]*)\]\(([^)]+)\)')
TABLE_SEPARATOR_RE = re.compile(r'^\s*\|?\s*:?-{3,}:?\s*(?:\|\s*:?-{3,}:?\s*)+\|?\s*$')
MATH_ENV_NAMES = r'equation\*?|align\*?|alignat\*?|aligned|gather\*?|multline\*?|flalign\*?'
DISPLAY_MATH_ENV_START_RE = re.compile(rf'^\s*\\begin\{{({MATH_ENV_NAMES})\}}')
DISPLAY_MATH_ENV_END_RE = re.compile(r'\\end\{([^}]+)\}')
PANDOC_MATH_PROTECTED_RE = re.compile(
    r'(```[\s\S]*?```|`[^`\n]*`|!\[[^\]]*\]\([^)]+\)|\[[^\]]+\]\([^)]+\)|\$\$[\s\S]*?\$\$|(?<!\\)\$[^\n$]+?(?<!\\)\$)'
)
BARE_LATEX_TOKEN_RE = re.compile(
    r'(?<![$\\\w./-])((?:\\[A-Za-z]+|[A-Za-z][A-Za-z0-9]*)(?:\s*(?:[_^]\{[^{}\n]+\}|[_^][A-Za-z0-9]+))+)(?![$\w./-])'
)


def export_paper_document(payload, project_dir, web_dir):
    fmt = str(payload.get('format', 'docx') or 'docx').strip().lower()
    if fmt not in {'docx', 'md', 'markdown', 'txt'}:
        raise ValueError('暂不支持该导出格式')

    sections = _normalize_sections(payload)
    fallback_text = str(payload.get('content', '') or payload.get('text', '') or '').strip()
    if not sections and fallback_text:
        sections = [{'title': '正文', 'display_title': '正文', 'level': 1, 'kind': '', 'content': fallback_text}]
    if not sections:
        raise ValueError('没有可导出的论文内容')

    title = _clean_text(payload.get('title') or '论文导出')
    base_name = _safe_filename(title or '论文导出')
    if fmt == 'docx':
        body = _build_docx(payload, sections, project_dir, web_dir)
        return ExportedPaper(
            filename=f'{base_name}.docx',
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            body=body,
        )
    if fmt in {'md', 'markdown'}:
        return ExportedPaper(
            filename=f'{base_name}.md',
            content_type='text/markdown; charset=utf-8',
            body=_build_markdown(payload, sections).encode('utf-8'),
        )
    return ExportedPaper(
        filename=f'{base_name}.txt',
        content_type='text/plain; charset=utf-8',
        body=_build_text(payload, sections).encode('utf-8'),
    )


def _normalize_sections(payload):
    result = []
    for item in payload.get('sections') or []:
        if not isinstance(item, dict):
            continue
        title = _clean_text(item.get('title') or '')
        display_title = _clean_text(item.get('displayTitle') or item.get('display_title') or title)
        content = _normalize_newlines(item.get('content') or '').strip()
        has_children = bool(item.get('hasChildren') or item.get('has_children'))
        if not title and not display_title and not content:
            continue
        if not content and not has_children:
            continue
        try:
            level = int(item.get('level') or 1)
        except Exception:
            level = 1
        result.append({
            'title': title or display_title or '未命名章节',
            'display_title': display_title or title or '未命名章节',
            'level': max(1, min(level, 4)),
            'kind': str(item.get('kind', '') or '').strip(),
            'content': content,
            'has_children': has_children,
        })
    return result


def _build_markdown(payload, sections):
    parts = []
    title = _clean_text(payload.get('title') or '')
    if title:
        parts.append(f'# {title}')
    for section in sections:
        heading_level = max(1, min(section['level'], 4))
        if section['display_title']:
            parts.append(f'{"#" * heading_level} {section["display_title"]}')
        if section['content']:
            parts.append(section['content'])
    return '\n\n'.join(part.strip() for part in parts if str(part or '').strip()).strip() + '\n'


def _build_text(payload, sections):
    markdown = _build_markdown(payload, sections)
    text = re.sub(r'^\s*#{1,6}\s+', '', markdown, flags=re.M)
    text = re.sub(IMAGE_LINE_RE, lambda match: f'[图片：{match.group(1) or "论文图表"}] {match.group(2)}', text)
    text = re.sub(r'\[([^\]]+)\]\(([^)]+)\)', r'\1（\2）', text)
    text = re.sub(r'[*_`~]+', '', text)
    return text.strip() + '\n'


def _build_docx(payload, sections, project_dir, web_dir):
    pandoc = _find_pandoc()
    if pandoc:
        try:
            return _build_docx_with_pandoc(payload, sections, project_dir, web_dir, pandoc)
        except Exception:
            pass
    return _build_docx_with_python_docx(payload, sections, project_dir, web_dir)


def _build_docx_with_pandoc(payload, sections, project_dir, web_dir, pandoc):
    with tempfile.TemporaryDirectory(prefix='paper_export_') as temp_dir:
        markdown = _build_markdown(payload, sections)
        markdown = _normalize_export_markdown_blocks(markdown)
        markdown = _normalize_math_for_pandoc(markdown)
        markdown = _localize_markdown_images(markdown, temp_dir, project_dir, web_dir)
        input_path = os.path.join(temp_dir, 'paper.md')
        output_path = os.path.join(temp_dir, 'paper.docx')
        with open(input_path, 'w', encoding='utf-8', newline='\n') as handle:
            handle.write(markdown)
        resource_path = os.pathsep.join([temp_dir, web_dir, project_dir])
        subprocess.run(
            [
                pandoc,
                input_path,
                '-f',
                'markdown+tex_math_dollars+tex_math_single_backslash+pipe_tables',
                '-t',
                'docx',
                '--standalone',
                '--resource-path',
                resource_path,
                '-o',
                output_path,
            ],
            cwd=temp_dir,
            check=True,
            stdout=subprocess.DEVNULL,
            stderr=subprocess.DEVNULL,
            timeout=60,
        )
        _polish_docx_tables(output_path)
        with open(output_path, 'rb') as handle:
            return handle.read()


def _build_docx_with_python_docx(payload, sections, project_dir, web_dir):
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.shared import Inches, Pt

    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.1)
    section.right_margin = Inches(1.1)
    _configure_docx_styles(document, qn, Pt)

    title = _clean_text(payload.get('title') or '')
    if title:
        title_para = document.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title_para.add_run(title)
        run.bold = True
        run.font.size = Pt(18)
        _set_run_font(run, qn, east_asia='黑体', ascii_font='Times New Roman')

    meta_items = [
        _clean_text(payload.get('subject') or ''),
        _clean_text(payload.get('paperStyle') or payload.get('paper_style') or ''),
        _clean_text(payload.get('referenceStyle') or payload.get('reference_style') or ''),
    ]
    meta_text = ' / '.join(item for item in meta_items if item)
    if meta_text:
        meta_para = document.add_paragraph()
        meta_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = meta_para.add_run(meta_text)
        run.font.size = Pt(10)
        _set_run_font(run, qn)

    if title or meta_text:
        document.add_paragraph()

    for item in sections:
        heading_text = item['display_title']
        if heading_text:
            heading = document.add_heading(heading_text, level=max(1, min(item['level'], 3)))
            for run in heading.runs:
                _set_run_font(run, qn, east_asia='黑体', ascii_font='Times New Roman')
        if item['content']:
            _add_markdown_content(document, item['content'], item.get('kind', ''), project_dir, web_dir, qn, Inches, Pt, WD_ALIGN_PARAGRAPH)

    stream = io.BytesIO()
    document.save(stream)
    return stream.getvalue()


def _configure_docx_styles(document, qn, Pt):
    normal = document.styles['Normal']
    normal.font.name = 'Times New Roman'
    normal._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    normal.font.size = Pt(12)
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.space_after = Pt(6)
    for style_name, size in [('Heading 1', 16), ('Heading 2', 14), ('Heading 3', 13)]:
        style = document.styles[style_name]
        style.font.name = 'Times New Roman'
        style._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        style.font.size = Pt(size)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(12)
        style.paragraph_format.space_after = Pt(6)


def _add_markdown_content(document, content, kind, project_dir, web_dir, qn, Inches, Pt, WD_ALIGN_PARAGRAPH):
    lines = _normalize_export_markdown_blocks(content).split('\n')
    index = 0
    is_reference = kind == 'reference'
    while index < len(lines):
        line = lines[index].rstrip()
        if not line.strip():
            index += 1
            continue

        math_block = _collect_display_math(lines, index)
        if math_block:
            formula, consumed = math_block
            _add_docx_formula_block(document, formula, qn, Pt, WD_ALIGN_PARAGRAPH)
            index += consumed
            continue

        image = _parse_image_line(line)
        if image:
            _add_docx_image(document, image['src'], image['alt'], project_dir, web_dir, Inches, WD_ALIGN_PARAGRAPH)
            index += 1
            continue

        table_block = _collect_markdown_table(lines, index)
        if table_block:
            caption_lines, table_rows, note_lines, consumed = table_block
            if caption_lines:
                _add_docx_text_block(document, caption_lines, False, qn, Pt, WD_ALIGN_PARAGRAPH)
            _add_docx_table(document, table_rows, qn, Pt)
            if note_lines:
                _add_docx_text_block(document, note_lines, False, qn, Pt, WD_ALIGN_PARAGRAPH)
            index += consumed
            continue

        heading = re.match(r'^\s*(#{1,6})\s+(.+?)\s*$', line)
        if heading:
            document.add_heading(_strip_inline_markdown(heading.group(2)), level=min(len(heading.group(1)), 3))
            index += 1
            continue

        block = [line]
        index += 1
        while index < len(lines) and lines[index].strip() and not _parse_image_line(lines[index]) and not _collect_markdown_table(lines, index):
            block.append(lines[index].rstrip())
            index += 1
        _add_docx_text_block(document, block, is_reference, qn, Pt, WD_ALIGN_PARAGRAPH)


def _add_docx_text_block(document, lines, is_reference, qn, Pt, WD_ALIGN_PARAGRAPH):
    for line in lines:
        stripped = line.strip()
        if not stripped:
            continue
        bullet = re.match(r'^\s*[-*•]\s+(.+)$', stripped)
        numbered = re.match(r'^\s*\d+[.)、]\s+(.+)$', stripped)
        caption = re.match(r'^\s*((?:图|表)\s*\d+(?:\.\d+)?|Figure\s*\d+)', stripped, flags=re.I)
        table_unit = re.match(r'^\s*[（(]单位[:：].+[）)]\s*$', stripped)
        table_note = re.match(r'^\s*(?:资料来源|注[:：]|备注[:：])', stripped)
        if bullet:
            paragraph = document.add_paragraph(style='List Bullet')
            paragraph.add_run(_strip_inline_markdown(bullet.group(1)))
        elif numbered:
            paragraph = document.add_paragraph(style='List Number')
            paragraph.add_run(_strip_inline_markdown(numbered.group(1)))
        else:
            paragraph = document.add_paragraph()
            paragraph.add_run(_strip_inline_markdown(stripped))
        if caption or table_unit:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if table_note:
            paragraph.paragraph_format.first_line_indent = Pt(0)
            for run in paragraph.runs:
                run.font.size = Pt(9)
        if is_reference or re.match(r'^\s*\[\d+\]', stripped):
            paragraph.paragraph_format.left_indent = Pt(21)
            paragraph.paragraph_format.first_line_indent = Pt(-21)
            paragraph.paragraph_format.line_spacing = 1.0
        for run in paragraph.runs:
            _set_run_font(run, qn)


def _collect_display_math(lines, index):
    line = lines[index].strip()
    if line.startswith('$$'):
        collected = [line]
        cursor = index + 1
        if line.count('$$') >= 2 and len(line) > 2:
            return _strip_math_delimiters('\n'.join(collected)), 1
        while cursor < len(lines):
            collected.append(lines[cursor].strip())
            if '$$' in lines[cursor]:
                break
            cursor += 1
        return _strip_math_delimiters('\n'.join(collected)), max(1, cursor - index + 1)
    if line.startswith('\\['):
        collected = [line]
        cursor = index + 1
        if '\\]' in line:
            return _strip_math_delimiters('\n'.join(collected)), 1
        while cursor < len(lines):
            collected.append(lines[cursor].strip())
            if '\\]' in lines[cursor]:
                break
            cursor += 1
        return _strip_math_delimiters('\n'.join(collected)), max(1, cursor - index + 1)
    env_match = DISPLAY_MATH_ENV_START_RE.match(line)
    if env_match:
        env_name = env_match.group(1)
        collected = [line]
        cursor = index + 1
        if re.search(rf'\\end\{{{re.escape(env_name)}\}}', line):
            return '\n'.join(collected).strip(), 1
        while cursor < len(lines):
            collected.append(lines[cursor].strip())
            if re.search(rf'\\end\{{{re.escape(env_name)}\}}', lines[cursor]):
                break
            cursor += 1
        return '\n'.join(collected).strip(), max(1, cursor - index + 1)
    return None


def _add_docx_formula_block(document, formula, qn, Pt, WD_ALIGN_PARAGRAPH):
    image = _render_formula_png(formula)
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if image:
        try:
            paragraph.add_run().add_picture(image)
            return
        except Exception:
            pass
    run = paragraph.add_run(_strip_math_delimiters(formula))
    run.font.name = 'Cambria Math'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Cambria Math')
    run.font.size = Pt(12)


def _render_formula_png(formula):
    try:
        from matplotlib import mathtext
        from matplotlib.font_manager import FontProperties
    except Exception:
        return None
    expression = _mathtext_expression(formula)
    if not expression:
        return None
    stream = io.BytesIO()
    try:
        mathtext.math_to_image(expression, stream, prop=FontProperties(size=14), dpi=220, format='png')
    except Exception:
        return None
    stream.seek(0)
    return stream


def _mathtext_expression(formula):
    text = _strip_math_delimiters(formula)
    if not text:
        return ''
    if '\\begin{' in text:
        return ''
    return f'${text}$'


def _normalize_export_markdown_blocks(content):
    flattened_lines = []
    for line in _normalize_newlines(content).split('\n'):
        flattened_lines.extend(_normalize_flat_pipe_table_line(line))
    return '\n'.join(_separate_markdown_tables(flattened_lines))


def _separate_markdown_tables(lines):
    output = []
    index = 0
    while index < len(lines):
        line = lines[index]
        if index + 1 < len(lines) and '|' in line and TABLE_SEPARATOR_RE.match(lines[index + 1]):
            if output and output[-1].strip():
                output.append('')
            while index < len(lines) and lines[index].strip() and (
                '|' in lines[index] or TABLE_SEPARATOR_RE.match(lines[index])
            ):
                output.append(lines[index])
                index += 1
            if index < len(lines) and lines[index].strip():
                output.append('')
            continue
        output.append(line)
        index += 1
    return output


def _normalize_flat_pipe_table_line(line):
    text = str(line or '').rstrip()
    if '|' not in text:
        return [text]
    separator_match = re.search(r'\|\s*:?-{3,}:?\s*(?:\|\s*:?-{3,}:?\s*)+\|', text)
    if not separator_match:
        return [text]
    separator = separator_match.group(0).strip()
    column_count = len(_split_table_row(separator))
    if column_count < 2:
        return [text]

    before = text[:separator_match.start()].rstrip()
    pipe_positions = [match.start() for match in re.finditer(r'\|', before)]
    if len(pipe_positions) < column_count + 1:
        return [text]
    header_start = pipe_positions[-(column_count + 1)]
    prefix = before[:header_start].strip()
    header = before[header_start:].strip()
    if not header.startswith('|'):
        return [text]

    output = []
    if prefix:
        caption_match = re.search(r'((?:图|表)\s*\d+(?:\.\d+)?[^\n|]*)$', prefix)
        if caption_match:
            preface = prefix[:caption_match.start()].strip()
            caption = caption_match.group(1).strip()
            if preface:
                output.append(preface)
            output.append(caption)
        else:
            output.append(prefix)
    output.extend([header, separator])

    cursor = separator_match.end()
    while cursor < len(text):
        while cursor < len(text) and text[cursor].isspace():
            cursor += 1
        if cursor >= len(text):
            break
        if text[cursor] != '|':
            tail = text[cursor:].strip()
            if tail:
                output.append(tail)
            break
        pipe_count = 0
        end = cursor
        while end < len(text):
            if text[end] == '|':
                pipe_count += 1
                if pipe_count == column_count + 1:
                    end += 1
                    break
            end += 1
        if pipe_count < column_count + 1:
            tail = text[cursor:].strip()
            if tail:
                output.append(tail)
            break
        output.append(text[cursor:end].strip())
        cursor = end

    return output or [text]


def _collect_markdown_table(lines, index):
    if index + 1 >= len(lines):
        return None
    caption_lines = []
    start = index
    for offset in range(0, 3):
        cursor = index + offset
        if cursor + 1 >= len(lines):
            break
        header = lines[cursor]
        separator = lines[cursor + 1]
        if '|' in header and TABLE_SEPARATOR_RE.match(separator):
            start = cursor
            caption_lines = [line.strip() for line in lines[index:start] if line.strip()]
            break
    else:
        return None

    header = lines[start]
    separator = lines[start + 1]
    if '|' not in header or not TABLE_SEPARATOR_RE.match(separator):
        return None
    rows = [_split_table_row(header)]
    cursor = start + 2
    while cursor < len(lines) and '|' in lines[cursor] and lines[cursor].strip():
        rows.append(_split_table_row(lines[cursor]))
        cursor += 1
    note_lines = []
    while cursor < len(lines) and lines[cursor].strip() and re.match(r'^\s*(?:资料来源|注[:：]|备注[:：])', lines[cursor].strip()):
        note_lines.append(lines[cursor].strip())
        cursor += 1
    return caption_lines, rows, note_lines, cursor - index


def _split_table_row(line):
    text = line.strip().strip('|')
    return [_strip_inline_markdown(cell.strip()) for cell in text.split('|')]


def _add_docx_table(document, rows, qn, Pt):
    if not rows:
        return
    width = max(len(row) for row in rows)
    table = document.add_table(rows=len(rows), cols=width)
    _style_docx_table(table, qn, Pt)
    for row_index, row in enumerate(rows):
        for col_index in range(width):
            cell = table.cell(row_index, col_index)
            cell.text = row[col_index] if col_index < len(row) else ''
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 1.0
                for run in paragraph.runs:
                    _set_run_font(run, qn)
                    run.font.size = Pt(10.5)
                    if row_index == 0:
                        run.bold = True


def _polish_docx_tables(path):
    try:
        from docx import Document
        from docx.oxml.ns import qn
        from docx.shared import Pt
    except Exception:
        return
    try:
        document = Document(path)
    except Exception:
        return
    changed = False
    for table in document.tables:
        _style_docx_table(table, qn, Pt)
        for row_index, row in enumerate(table.rows):
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.space_before = Pt(0)
                    paragraph.paragraph_format.space_after = Pt(0)
                    paragraph.paragraph_format.line_spacing = 1.0
                    for run in paragraph.runs:
                        _set_run_font(run, qn)
                        run.font.size = Pt(10.5)
                        if row_index == 0:
                            run.bold = True
        changed = True
    if changed:
        document.save(path)


def _style_docx_table(table, qn, Pt):
    try:
        from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
        from docx.oxml import OxmlElement
    except Exception:
        WD_CELL_VERTICAL_ALIGNMENT = None
        WD_TABLE_ALIGNMENT = None
        OxmlElement = None
    try:
        table.style = 'Table Grid'
    except Exception:
        pass
    try:
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
    except Exception:
            pass
    try:
        table.autofit = False
        table.allow_autofit = False
    except Exception:
        pass
    _set_table_layout(table, qn, OxmlElement)
    column_count = max((len(row.cells) for row in table.rows), default=1)
    cell_width = max(900, int(8640 / max(1, column_count)))
    for row in table.rows:
        for cell in row.cells:
            if WD_CELL_VERTICAL_ALIGNMENT is not None:
                try:
                    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                except Exception:
                    pass
            _set_cell_width(cell, qn, OxmlElement, cell_width)
            _set_cell_margins(cell, qn, OxmlElement, top=60, bottom=60, left=90, right=90)
    _apply_three_line_table_borders(table, qn, OxmlElement)


def _set_table_layout(table, qn, OxmlElement, width=8640):
    if OxmlElement is None:
        return
    tbl_pr = table._tbl.tblPr
    if tbl_pr is None:
        tbl_pr = OxmlElement('w:tblPr')
        table._tbl.insert(0, tbl_pr)
    tbl_w = tbl_pr.find(qn('w:tblW'))
    if tbl_w is None:
        tbl_w = OxmlElement('w:tblW')
        tbl_pr.append(tbl_w)
    tbl_w.set(qn('w:w'), str(width))
    tbl_w.set(qn('w:type'), 'dxa')

    tbl_layout = tbl_pr.find(qn('w:tblLayout'))
    if tbl_layout is None:
        tbl_layout = OxmlElement('w:tblLayout')
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn('w:type'), 'fixed')

    borders = tbl_pr.find(qn('w:tblBorders'))
    if borders is None:
        borders = OxmlElement('w:tblBorders')
        tbl_pr.append(borders)
    for border_name in ('top', 'bottom'):
        _set_border_node(borders, border_name, qn, OxmlElement, val='single', size='12')
    for border_name in ('left', 'right', 'insideV'):
        _set_border_node(borders, border_name, qn, OxmlElement, val='nil', size='0')
    _set_border_node(borders, 'insideH', qn, OxmlElement, val='single', size='4', color='D9D9D9')


def _set_border_node(parent, border_name, qn, OxmlElement, *, val='single', size='4', color='auto'):
    if OxmlElement is None:
        return
    node = parent.find(qn(f'w:{border_name}'))
    if node is None:
        node = OxmlElement(f'w:{border_name}')
        parent.append(node)
    node.set(qn('w:val'), val)
    node.set(qn('w:sz'), str(size))
    node.set(qn('w:space'), '0')
    node.set(qn('w:color'), color)


def _apply_three_line_table_borders(table, qn, OxmlElement):
    if OxmlElement is None or not getattr(table, 'rows', None):
        return
    for row_index, row in enumerate(table.rows):
        for cell in row.cells:
            tc_pr = cell._tc.get_or_add_tcPr()
            borders = tc_pr.find(qn('w:tcBorders'))
            if borders is None:
                borders = OxmlElement('w:tcBorders')
                tc_pr.append(borders)
            _set_border_node(borders, 'left', qn, OxmlElement, val='nil', size='0')
            _set_border_node(borders, 'right', qn, OxmlElement, val='nil', size='0')
            _set_border_node(borders, 'top', qn, OxmlElement, val='single', size='4', color='D9D9D9')
            _set_border_node(borders, 'bottom', qn, OxmlElement, val='single', size='4', color='D9D9D9')
            if row_index == 0:
                _set_border_node(borders, 'top', qn, OxmlElement, val='single', size='12')
                _set_border_node(borders, 'bottom', qn, OxmlElement, val='single', size='8')
            if row_index == len(table.rows) - 1:
                _set_border_node(borders, 'bottom', qn, OxmlElement, val='single', size='12')


def _set_cell_width(cell, qn, OxmlElement, width):
    if OxmlElement is None:
        return
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn('w:tcW'))
    if tc_w is None:
        tc_w = OxmlElement('w:tcW')
        tc_pr.append(tc_w)
    tc_w.set(qn('w:w'), str(width))
    tc_w.set(qn('w:type'), 'dxa')


def _set_cell_margins(cell, qn, OxmlElement, top=60, bottom=60, left=90, right=90):
    if OxmlElement is None:
        return
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in('w:tcMar')
    if tc_mar is None:
        tc_mar = OxmlElement('w:tcMar')
        tc_pr.append(tc_mar)
    for margin_name, value in {
        'top': top,
        'bottom': bottom,
        'left': left,
        'right': right,
    }.items():
        node = tc_mar.find(qn(f'w:{margin_name}'))
        if node is None:
            node = OxmlElement(f'w:{margin_name}')
            tc_mar.append(node)
        node.set(qn('w:w'), str(value))
        node.set(qn('w:type'), 'dxa')


def _add_docx_image(document, src, alt, project_dir, web_dir, Inches, WD_ALIGN_PARAGRAPH):
    image_stream_or_path = _resolve_image_source(src, project_dir, web_dir)
    if not image_stream_or_path:
        paragraph = document.add_paragraph(f'[图片：{alt or "论文图表"}] {src}')
        return
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    try:
        run.add_picture(image_stream_or_path, width=Inches(5.7))
    except Exception:
        paragraph.add_run(f'[图片：{alt or "论文图表"}] {src}')


def _resolve_image_source(src, project_dir, web_dir):
    value = str(src or '').strip()
    if not value:
        return None
    if value.lower().startswith('data:image/') and ',' in value:
        try:
            raw = base64.b64decode(value.split(',', 1)[1], validate=False)
            return io.BytesIO(raw)
        except Exception:
            return None
    parsed = urlparse(value)
    if parsed.scheme in {'http', 'https'}:
        return None
    candidates = []
    if parsed.scheme == 'file':
        candidates.append(unquote(parsed.path))
    else:
        relative = unquote(value.lstrip('/\\'))
        candidates.append(os.path.join(web_dir, relative))
        candidates.append(os.path.join(project_dir, relative))
        if os.path.isabs(value):
            candidates.append(value)
    for candidate in candidates:
        path = os.path.abspath(candidate)
        if os.path.isfile(path):
            if path.startswith(os.path.abspath(project_dir)) or path.startswith(os.path.abspath(web_dir)):
                return path
    return None


def _find_pandoc():
    candidates = [shutil.which('pandoc')]
    for path in (
        r'D:\Anaconda3\Scripts\pandoc.exe',
        r'C:\ProgramData\Anaconda3\Scripts\pandoc.exe',
        r'C:\Users\Hang\Anaconda3\Scripts\pandoc.exe',
        r'C:\Program Files\Pandoc\pandoc.exe',
    ):
        candidates.append(path)
    for candidate in candidates:
        if candidate and os.path.isfile(candidate):
            return candidate
    return ''


def _normalize_math_for_pandoc(markdown):
    text = _normalize_newlines(markdown)
    text = re.sub(r'\\\[\s*([\s\S]*?)\s*\\\]', _replace_display_math_delimiters, text)
    text = re.sub(r'\\\(([\s\S]*?)\\\)', _replace_inline_math_delimiters, text)

    pattern = re.compile(rf'(^|\n)(\\begin\{{(?:{MATH_ENV_NAMES})\}}[\s\S]*?\\end\{{(?:{MATH_ENV_NAMES})\}})(?=\n|$)')

    def replace(match):
        prefix = match.group(1)
        block = match.group(2).strip()
        before = text[:match.start()]
        if before.rstrip().endswith('$$'):
            return match.group(0)
        return f'{prefix}$$\n{block}\n$$'

    text = pattern.sub(replace, text)
    text = _normalize_bare_display_math_lines(text)
    return _normalize_bare_inline_math(text)


def _replace_display_math_delimiters(match):
    formula = match.group(1).strip()
    if not formula:
        return ''
    return f'\n$$\n{formula}\n$$\n'


def _replace_inline_math_delimiters(match):
    formula = re.sub(r'\s+', ' ', match.group(1).strip())
    if not formula:
        return ''
    return f'${formula}$'


def _normalize_bare_display_math_lines(text):
    result = []
    in_code = False
    in_math = False
    for line in _normalize_newlines(text).split('\n'):
        stripped = line.strip()
        if stripped.startswith('```'):
            result.append(line)
            in_code = not in_code
            continue
        if not in_code and _has_unmatched_display_dollars(stripped):
            result.append(line)
            in_math = not in_math
            continue
        if not in_code and not in_math and _looks_like_standalone_math_line(stripped):
            indent = line[:len(line) - len(line.lstrip())]
            result.extend([f'{indent}$$', f'{indent}{stripped}', f'{indent}$$'])
            continue
        result.append(line)
    return '\n'.join(result)


def _has_unmatched_display_dollars(line):
    return line.count('$$') % 2 == 1


def _looks_like_standalone_math_line(line):
    if not line or len(line) > 320:
        return False
    if line.startswith(('#', '>', '|', '-', '*', '+', '!', '[', '$', '\\[', '\\(')):
        return False
    if re.search(r'[\u4e00-\u9fff]', line):
        return False
    if not _is_likely_math_expression(line):
        return False
    return bool(re.search(r'(?:=|≤|≥|≈|≠|<|>|\\frac|\\sum|\\prod|\\int|\\sqrt)', line))


def _normalize_bare_inline_math(text):
    parts = []
    cursor = 0
    for match in PANDOC_MATH_PROTECTED_RE.finditer(text):
        parts.append(_normalize_bare_inline_math_segment(text[cursor:match.start()]))
        parts.append(match.group(0))
        cursor = match.end()
    parts.append(_normalize_bare_inline_math_segment(text[cursor:]))
    return ''.join(parts)


def _normalize_bare_inline_math_segment(segment):
    segment = re.sub(r'([（(])\s*([^()（）\n]{1,180})\s*([）)])', _replace_parenthesized_math, segment)
    return BARE_LATEX_TOKEN_RE.sub(lambda match: f'${match.group(1).strip()}$', segment)


def _replace_parenthesized_math(match):
    opener = match.group(1)
    body = match.group(2).strip()
    closer = match.group(3)
    if not _is_likely_math_expression(body):
        return match.group(0)
    return f'{opener}${body}${closer}'


def _is_likely_math_expression(value):
    text = str(value or '').strip()
    if not text or len(text) > 220:
        return False
    if '://' in text or '@' in text:
        return False
    if re.match(r'^[A-Z][A-Za-z .-]+,\s*(?:19|20)\d{2}', text):
        return False
    if not re.search(r'[A-Za-zΑ-Ωα-ω\\]', text):
        return False
    if re.search(r'(?:[_^]\{[^{}]+\}|[_^][A-Za-z0-9]+|\\[A-Za-z]+)', text):
        return True
    if re.search(r'[=+*/<>≤≥≈≠±×÷∑∏√∞]', text):
        return True
    return bool(re.search(r'[Α-Ωα-ω]', text))


def _localize_markdown_images(markdown, temp_dir, project_dir, web_dir):
    image_counter = {'value': 0}

    def replace(match):
        alt = match.group(1)
        src = match.group(2).strip()
        localized = _local_image_path_for_pandoc(src, temp_dir, project_dir, web_dir, image_counter)
        if not localized:
            return match.group(0)
        return f'![{alt}]({localized})'

    return MARKDOWN_IMAGE_RE.sub(replace, markdown)


def _local_image_path_for_pandoc(src, temp_dir, project_dir, web_dir, image_counter):
    value = str(src or '').strip()
    if not value:
        return ''
    if value.lower().startswith('data:image/') and ',' in value:
        header, raw_data = value.split(',', 1)
        ext = 'png'
        match = re.match(r'data:image/([a-zA-Z0-9.+-]+)', header)
        if match:
            ext = {'jpeg': 'jpg', 'svg+xml': 'svg'}.get(match.group(1).lower(), match.group(1).lower())
        try:
            raw = base64.b64decode(raw_data, validate=False)
        except Exception:
            return ''
        image_counter['value'] += 1
        path = os.path.join(temp_dir, f'image-{image_counter["value"]}.{ext}')
        with open(path, 'wb') as handle:
            handle.write(raw)
        return path.replace('\\', '/')
    source = _resolve_image_source(value, project_dir, web_dir)
    if isinstance(source, str) and os.path.isfile(source):
        image_counter['value'] += 1
        _, ext = os.path.splitext(source)
        ext = ext if ext.lower() in {'.png', '.jpg', '.jpeg', '.gif', '.bmp', '.svg', '.webp'} else '.png'
        local_name = f'image-{image_counter["value"]}{ext}'
        local_path = os.path.join(temp_dir, local_name)
        shutil.copyfile(source, local_path)
        return local_name
    return ''


def _parse_image_line(line):
    match = IMAGE_LINE_RE.match(str(line or ''))
    if not match:
        return None
    return {'alt': match.group(1).strip(), 'src': match.group(2).strip()}


def _strip_inline_markdown(text):
    value = str(text or '')
    value = re.sub(r'!\[([^\]]*)\]\(([^)]+)\)', r'\1', value)
    value = re.sub(r'\[([^\]]+)\]\(([^)]+)\)', r'\1', value)
    value = re.sub(r'(\*\*|__)(.*?)\1', r'\2', value)
    value = re.sub(r'(\*|_)(.*?)\1', r'\2', value)
    value = re.sub(r'`([^`]+)`', r'\1', value)
    return value.strip()


def _strip_math_delimiters(text):
    value = _normalize_newlines(text).strip()
    if value.startswith('$$') and value.endswith('$$'):
        return value[2:-2].strip()
    if value.startswith('\\[') and value.endswith('\\]'):
        return value[2:-2].strip()
    if value.startswith('\\(') and value.endswith('\\)'):
        return value[2:-2].strip()
    if value.startswith('$') and value.endswith('$') and len(value) > 1:
        return value[1:-1].strip()
    return value


def _set_run_font(run, qn, east_asia='宋体', ascii_font='Times New Roman'):
    run.font.name = ascii_font
    run._element.rPr.rFonts.set(qn('w:eastAsia'), east_asia)


def _normalize_newlines(value):
    return str(value or '').replace('\r\n', '\n').replace('\r', '\n')


def _clean_text(value):
    return re.sub(r'\s+', ' ', str(value or '').strip())


def _safe_filename(value):
    text = re.sub(r'[<>:"/\\|?*\x00-\x1f]+', '_', str(value or '').strip())
    text = re.sub(r'\s+', ' ', text).strip(' ._')
    return text[:80] or '论文导出'
